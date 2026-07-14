"""
Generate VP-ready Hurix Talent Assessment QA Go-Live Checklist (DOCX).

Visual direction (Hurix brand, not templated cream/terracotta or broadsheet):
- Palette: Ink slate, cyan→violet brand accents, clean white paper
- Signature: horizontal cyan–violet rule under the title band
- Structure: executive brief → evidence → module checklists → sign-off
"""

from datetime import date
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Hurix_Talent_QA_GoLive_Checklist.docx"

# Brand tokens
INK = "0F172A"
SLATE = "334155"
MUTED = "64748B"
CYAN = "06B6D4"
BLUE = "3B82F6"
VIOLET = "8B5CF6"
PAPER = "FFFFFF"
ROW_ALT = "F1F5F9"
PASS_BG = "ECFDF5"
PASS_FG = "047857"
FAIL_BG = "FEF2F2"
FAIL_FG = "B91C1C"
WARN_BG = "FFFBEB"
WARN_FG = "B45309"
HEADER_BG = "0F172A"


def set_cell_shading(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def clear_paragraph_spacing(p) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_runs(paragraph, segments, *, size=10, bold=False, color=INK):
    """segments: list of (text, bold?, color?) or plain strings mixed via helper."""
    for seg in segments:
        if isinstance(seg, str):
            text, is_bold, col = seg, bold, color
        else:
            text = seg[0]
            is_bold = seg[1] if len(seg) > 1 else bold
            col = seg[2] if len(seg) > 2 else color
        run = paragraph.add_run(text)
        run.bold = is_bold
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(col)
        run.font.name = "Calibri"


def style_table_borders(table, color=CYAN, sz="8") -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    # remove existing borders if any
    for child in list(tblPr):
        if child.tag == qn("w:tblBorders"):
            tblPr.remove(child)
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def add_title_band(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, INK)
    set_cell_margins(cell, top=140, bottom=120, left=160, right=160)

    p1 = cell.paragraphs[0]
    clear_paragraph_spacing(p1)
    add_runs(
        p1,
        [("HURIX DIGITAL", True, CYAN)],
        size=9,
    )
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p2 = cell.add_paragraph()
    clear_paragraph_spacing(p2)
    p2.paragraph_format.space_before = Pt(4)
    add_runs(p2, [("Talent Assessment Platform", True, PAPER)], size=18)
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p3 = cell.add_paragraph()
    clear_paragraph_spacing(p3)
    p3.paragraph_format.space_before = Pt(2)
    add_runs(
        p3,
        [("QA Go-Live Checklist  ·  Pre-Production Release Gate", False, "94A3B8")],
        size=11,
    )

    # Cyan–violet signature rule
    rule = doc.add_table(rows=1, cols=2)
    rule.alignment = WD_TABLE_ALIGNMENT.CENTER
    c0, c1 = rule.cell(0, 0), rule.cell(0, 1)
    set_cell_shading(c0, CYAN)
    set_cell_shading(c1, VIOLET)
    for c in (c0, c1):
        c.paragraphs[0].clear()
        set_cell_margins(c, top=18, bottom=18, left=0, right=0)
    # shrink row height feel via empty runs
    for c in (c0, c1):
        run = c.paragraphs[0].add_run(" ")
        run.font.size = Pt(2)

    doc.add_paragraph()


def add_section_heading(doc: Document, number: str, title: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    left, right = table.cell(0, 0), table.cell(0, 1)
    left.width = Cm(0.35)
    right.width = Cm(16.5)
    set_cell_shading(left, CYAN)
    set_cell_shading(right, PAPER)
    set_cell_margins(left, top=40, bottom=40, left=0, right=0)
    set_cell_margins(right, top=40, bottom=40, left=100, right=40)
    left.paragraphs[0].clear()
    rp = right.paragraphs[0]
    clear_paragraph_spacing(rp)
    add_runs(rp, [(f"{number}  ", True, VIOLET), (title, True, INK)], size=12)
    # spacer
    sp = doc.add_paragraph()
    clear_paragraph_spacing(sp)
    sp.paragraph_format.space_after = Pt(6)


def add_body(doc: Document, text: str, *, size=10, color=SLATE, space_after=6) -> None:
    p = doc.add_paragraph()
    clear_paragraph_spacing(p)
    p.paragraph_format.space_after = Pt(space_after)
    add_runs(p, [(text, False, color)], size=size)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    style_table_borders(table, color="E2E8F0", sz="4")
    for i, (k, v) in enumerate(rows):
        ck, cv = table.cell(i, 0), table.cell(i, 1)
        set_cell_shading(ck, ROW_ALT if i % 2 == 0 else PAPER)
        set_cell_shading(cv, ROW_ALT if i % 2 == 0 else PAPER)
        set_cell_margins(ck)
        set_cell_margins(cv)
        pk, pv = ck.paragraphs[0], cv.paragraphs[0]
        clear_paragraph_spacing(pk)
        clear_paragraph_spacing(pv)
        add_runs(pk, [(k, True, SLATE)], size=9)
        add_runs(pv, [(v, False, INK)], size=9)
    sp = doc.add_paragraph()
    clear_paragraph_spacing(sp)
    sp.paragraph_format.space_after = Pt(8)


def add_checklist_table(doc, headers, rows, status_col=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    style_table_borders(table, color="CBD5E1", sz="4")

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, HEADER_BG)
        set_cell_margins(cell, top=50, bottom=50, left=70, right=70)
        p = cell.paragraphs[0]
        clear_paragraph_spacing(p)
        add_runs(p, [(h, True, PAPER)], size=8)

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            bg = ROW_ALT if i % 2 else PAPER
            fg = INK
            if status_col is not None and j == status_col:
                key = (val or "").strip().upper()
                if key in ("PASS", "PASSED", "OK"):
                    bg, fg = PASS_BG, PASS_FG
                elif key in ("FAIL", "FAILED"):
                    bg, fg = FAIL_BG, FAIL_FG
                elif key in ("PARTIAL", "WARN", "ATTENTION", "OPEN"):
                    bg, fg = WARN_BG, WARN_FG
                elif key in ("SKIP", "N/A", "NA", "BLOCKED"):
                    bg, fg = "F8FAFC", MUTED
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=45, bottom=45, left=70, right=70)
            p = cell.paragraphs[0]
            clear_paragraph_spacing(p)
            add_runs(p, [(val, j == status_col, fg if j == status_col else SLATE)], size=8)

    sp = doc.add_paragraph()
    clear_paragraph_spacing(sp)
    sp.paragraph_format.space_after = Pt(10)


def add_blank_sign(doc: Document, labels: list[str]) -> None:
    table = doc.add_table(rows=len(labels), cols=3)
    style_table_borders(table, color="E2E8F0", sz="4")
    headers = ["Role", "Name / Signature", "Date"]
    # rebuild as header + rows
    table = doc.add_table(rows=1 + len(labels), cols=3)
    style_table_borders(table, color="CBD5E1", sz="4")
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, HEADER_BG)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        clear_paragraph_spacing(p)
        add_runs(p, [(h, True, PAPER)], size=9)
    for i, label in enumerate(labels):
        for j, val in enumerate([label, "", ""]):
            cell = table.rows[i + 1].cells[j]
            set_cell_shading(cell, PAPER)
            set_cell_margins(cell, top=90, bottom=90, left=80, right=80)
            p = cell.paragraphs[0]
            clear_paragraph_spacing(p)
            add_runs(p, [(val, j == 0, SLATE)], size=9)


def build() -> Path:
    today = date.today().strftime("%d %B %Y")
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)

    add_title_band(doc)

    # 1. Document control
    add_section_heading(doc, "01", "Document Control")
    add_kv_table(
        doc,
        [
            ("Document title", "QA Go-Live Checklist — Hurix Talent Assessment Platform"),
            ("Version", "1.0"),
            ("Prepared for", "Vice President — Product / Delivery sign-off"),
            ("Prepared by", "Engineering / QA (Product Development)"),
            ("Review date", today),
            ("Target release", "Admin portal candidate-management enhancements (LP → production)"),
            ("Environments covered", "Localhost (dev)  ·  EC2 dedicated test DB `hurix_talent_test` + Supabase from EC2 `.env`"),
            ("Classification", "Internal — pre-go-live gate"),
        ],
    )

    # 2. Executive summary
    add_section_heading(doc, "02", "Executive Summary")
    add_body(
        doc,
        "This checklist is the release gate for making the Talent Assessment admin portal "
        "live with the latest candidate-management capabilities: bulk outreach, templates, "
        "export, shortlist pipelines, filter hygiene, and non-blocking email send progress. "
        "It combines automated suite evidence from this QA cycle with manual verification "
        "items that require a human pass on staging/production before go-live.",
    )
    add_body(
        doc,
        "Recommendation from this cycle: CONDITIONAL GO — core admin journeys are functionally "
        "ready; RBAC/security automation passed 16/16 on EC2 against dedicated test DB "
        "`hurix_talent_test` (Supabase URL/keys from EC2 `.env`). Production DB `hurix_talent` "
        "was not used for destructive tests. Remaining stale non-security unit/integration "
        "assertions should be waived or updated before cutover.",
        color=WARN_FG,
    )

    # 3. Scope
    add_section_heading(doc, "03", "Release Scope Under Review")
    add_checklist_table(
        doc,
        ["#", "Capability", "Surface", "In scope?"],
        [
            ["1", "Candidate list filters, pagination, seamless page change", "Admin → Candidates", "Yes"],
            ["2", "Bulk actions: reminder, reject, shortlist, role, delete, export", "Admin → Candidates", "Yes"],
            ["3", "Non-blocking reminder send + progress overlay + error log", "Admin → Candidates", "Yes"],
            ["4", "Email / WhatsApp templates (plain-text email body → HTML)", "Admin → Templates", "Yes"],
            ["5", "Default sign-off text: Team Hurix Digital (no companyName variable)", "Templates / Reminders", "Yes"],
            ["6", "Export CSV / XLSX (all / filtered / selected)", "Admin → Export", "Yes"],
            ["7", "Shortlisted / Added / Rejected / Deleted list pages", "Admin navigation", "Yes"],
            ["8", "WhatsApp template picker → prefilled wa.me message", "Candidates / Shortlisted", "Yes"],
            ["9", "Score NA / Role unassigned (NA) filters", "Candidates / Shortlisted", "Yes"],
            ["10", "Candidate portal registration & assessment journey", "Public / Candidate portal", "Regression"],
            ["11", "Admin RBAC & security suites", "API", "Yes — PASS on EC2 test DB"],
        ],
    )

    # 4. Automated evidence
    add_section_heading(doc, "04", "Automated Test Evidence (This QA Cycle)")
    add_body(
        doc,
        f"Executed on {today}. Local: unit/frontend + API health. "
        "EC2: dedicated Postgres `hurix_talent_test` (isolated from prod `hurix_talent`) "
        "+ Supabase from EC2 `.env`; migrate + seed; then `npm run test:security` / related suites.",
    )
    add_checklist_table(
        doc,
        ["Suite", "Result", "Detail", "Go-live impact"],
        [
            [
                "Test database (EC2)",
                "PASS",
                "Used `hurix_talent_test` only. Prod `hurix_talent` not wiped. Seed applied (SUPER_ADMIN + questions).",
                "Safe, repeatable QA pattern on EC2 Postgres.",
            ],
            [
                "Backend security (EC2 + test DB)",
                "PASS",
                "16/16 passed — assessment token middleware + RBAC (dashboard, analytics, /admin/me, JWT isolation).",
                "Security gate cleared for this cycle.",
            ],
            [
                "Backend unit",
                "PARTIAL",
                "95 passed · 4 failed · 31 skipped (DB-dependent)",
                "Failures are assertion drift vs product intent (see §05). Not runtime blockers by themselves.",
            ],
            [
                "Frontend unit",
                "PARTIAL",
                "54 passed · 1 failed (LandingPage expects Candidate Login CTA moved/removed)",
                "Update test or restore CTA; low production risk if intentional UX change.",
            ],
            [
                "API smoke (local)",
                "PASS",
                "Hurix Talent API /api/health returned status=ok; SMTP verify logged on boot",
                "Local stack healthy for manual QA.",
            ],
            [
                "CI quality gates",
                "OPEN",
                "Refer docs/TESTING.md — GitHub Actions on main",
                "Confirm latest main pipeline green before production deploy.",
            ],
        ],
        status_col=1,
    )

    # 5. Known failures / waivers
    add_section_heading(doc, "05", "Known Automated Failures & Waiver Notes")
    add_checklist_table(
        doc,
        ["ID", "Failing check", "Observed cause", "Disposition"],
        [
            [
                "U-01",
                "experience.test labels / years",
                "Product labels shortened (e.g. Fresher); FIVE_SEVEN years mapped to 7",
                "Update unit expectations — product change, not defect",
            ],
            [
                "U-02",
                "candidate-selection excludes REJECTED",
                "Main list now excludes REJECTED + SHORTLISTED (pipeline split)",
                "Update unit expectations to notIn [REJECTED, SHORTLISTED]",
            ],
            [
                "U-03",
                "inactivityDays=14 validation throw",
                "Parser now accepts additional day values",
                "Align test with allowed values or document whitelist",
            ],
            [
                "F-01",
                "LandingPage Candidate Login link",
                "Hero CTA set is Apply Now only (login may live in header)",
                "Update RTL test to match current landing design",
            ],
        ],
    )

    # 6. Manual QA checklists
    add_section_heading(doc, "06", "Manual QA Checklist — Admin Candidate Management")
    add_body(doc, "Mark each row Pass / Fail / N/A during staging walkthrough. Tester initials optional in Notes.")
    add_checklist_table(
        doc,
        ["ID", "Test case", "Expected result", "Status", "Notes"],
        [
            ["A-01", "Admin login (ADMIN / SUPER_ADMIN)", "Session established; correct nav by role", "", ""],
            ["A-02", "Candidates list loads with total count", "Count accurate; no layout jump on pagination", "", ""],
            ["A-03", "Pagination arrows (prev/next)", "Seamless swap of rows; no Updating… text; no UI shift", "", ""],
            ["A-04", "Page size change (10/25/50…)", "List refreshes; header count remains stable", "", ""],
            ["A-05", "Search by name / email / application id", "Results filter correctly within ~350ms debounce", "", ""],
            ["A-06", "Experience / country multi-select filters", "AND behaviour with search; empty state clear", "", ""],
            ["A-07", "Score filter including NA (no submissions)", "Only candidates without scores when NA selected", "", ""],
            ["A-08", "Role assignment filter including NA/unassigned", "Unassigned candidates appear correctly", "", ""],
            ["A-09", "Registered date range calendar", "From/To apply; localhost light calendar if applicable", "", ""],
            ["A-10", "Role quick-filter chips", "Chip filter matches role column; All resets", "", ""],
            ["A-11", "Row select + select-all on page", "Bulk bar enables; effective count correct", "", ""],
            ["A-12", "Select all matching filter set", "Effective count = filtered total minus exclusions", "", ""],
            ["A-13", "Bulk Shortlist", "Selected move to Shortlisted; leave main list", "", ""],
            ["A-14", "Bulk Reject", "Selected appear under Rejected with reason if captured", "", ""],
            ["A-15", "Bulk Assign Role", "Role chip updates; filter reflects new role", "", ""],
            ["A-16", "Bulk Soft Delete", "Candidates leave active list; appear in Deleted", "", ""],
            ["A-17", "Add Candidate form", "Creates ADMIN_CREATED profile; visible in Added", "", ""],
            ["A-18", "Candidate detail / resume preview", "Profile fields + PDF preview open without error", "", ""],
            ["A-19", "Score breakdown drawer", "Marks per question; NA handled gracefully", "", ""],
            ["A-20", "Owner assignment (if enabled)", "Owner persists; list reflects owner filter", "", ""],
        ],
        status_col=3,
    )

    add_section_heading(doc, "07", "Manual QA Checklist — Outreach, Templates & Export")
    add_checklist_table(
        doc,
        ["ID", "Test case", "Expected result", "Status", "Notes"],
        [
            ["O-01", "Send Reminder — template picker + preview", "Preview subject/body; Cancel closes modal", "", ""],
            ["O-02", "Send Reminder — non-blocking progress", "Modal closes; floating progress under header; list UI does not shift", "", ""],
            ["O-03", "Progress counter advances (1 of N …)", "IDS selection shows true per-mail progress", "", ""],
            ["O-04", "Partial failure / missing email", "Error log expands with candidate references", "", ""],
            ["O-05", "Stuck send (>45s idle)", "Still working… affordance appears without locking UI", "", ""],
            ["O-06", "Email arrives (SMTP)", "Inbox receives rendered HTML; portal link present", "", ""],
            ["O-07", "Sign-off text", "Ends with Team Hurix Digital (not {{companyName}})", "", ""],
            ["O-08", "Templates — create email (plain text)", "Enter/paragraphs become HTML; variables preserved", "", ""],
            ["O-09", "Templates — edit existing email", "HTML opens as readable plain text; save re-stores HTML", "", ""],
            ["O-10", "Templates — WhatsApp create/edit/delete", "CRUD works; message body variables listed", "", ""],
            ["O-11", "WhatsApp icon on candidate", "Template modal → opens wa.me with filled text", "", ""],
            ["O-12", "Export CSV — All active", "File downloads; columns intact", "", ""],
            ["O-13", "Export XLSX — Filtered / Selected", "Disabled states correct when no filter/selection", "", ""],
            ["O-14", "Shortlisted page parity", "Filters, WhatsApp, reminder behaviours match Candidates", "", ""],
            ["O-15", "RBAC: non–Super Admin export", "Export hidden/denied per permission matrix", "", ""],
        ],
        status_col=3,
    )

    add_section_heading(doc, "08", "Manual QA Checklist — Candidate Portal Regression")
    add_checklist_table(
        doc,
        ["ID", "Test case", "Expected result", "Status", "Notes"],
        [
            ["C-01", "Landing page / Apply Now", "Hero renders; CTA to /register", "", ""],
            ["C-02", "Registration + resume upload", "Account created; verification email sent", "", ""],
            ["C-03", "Email verification link", "Marks verified; dashboard access", "", ""],
            ["C-04", "Candidate login", "Dashboard loads with timeline status", "", ""],
            ["C-05", "Start / continue assessment", "Timer, navigation, submit/timeout path", "", ""],
            ["C-06", "Mobile assessment block", "Mobile blocked with clear message", "", ""],
            ["C-07", "Portal link in reminder email", "https://candidates.hurixsystems.com/ reachable", "", ""],
        ],
        status_col=3,
    )

    add_section_heading(doc, "09", "Non-Functional & Production Readiness")
    add_checklist_table(
        doc,
        ["ID", "Check", "Expected", "Status", "Notes"],
        [
            ["N-01", "Environment variables on EC2", "DB, JWT, SMTP, frontend URL, CORS correct", "", ""],
            ["N-02", "Docker compose prod rebuild", "Containers healthy after deploy", "", ""],
            ["N-03", "HTTPS / reverse proxy", "Admin & candidate URLs secure where required", "", ""],
            ["N-04", "SMTP deliverability", "Reminder emails not junked for sample domains", "", ""],
            ["N-05", "Backup / rollback plan", "Prior image/tag noted; DB migrate reversible or backed up", "", ""],
            ["N-06", "PII handling", "Exports limited to entitled roles; audit log for bulk actions", "", ""],
            ["N-07", "Rate limits on bulk / auth", "Abuse paths return 429 without crash", "", ""],
            ["N-08", "Browser matrix", "Chrome + Edge latest; mobile admin usable for list", "", ""],
        ],
        status_col=3,
    )

    add_section_heading(doc, "10", "Defect Log (Open at Time of Issue)")
    add_checklist_table(
        doc,
        ["#", "Severity", "Description", "Owner", "Status"],
        [
            ["D-01", "Low", "Unit tests lagging product filter/experience semantics (§05)", "Eng", "OPEN"],
            ["D-02", "Low", "LandingPage RTL assertion outdated for Candidate Login CTA", "Eng", "OPEN"],
            ["D-03", "Low", "Security suite now PASS on EC2 `hurix_talent_test` (was blocked without test DB)", "Eng", "CLOSED"],
            ["", "", "", "", ""],
            ["", "", "", "", ""],
        ],
        status_col=4,
    )

    add_section_heading(doc, "11", "Go-Live Decision Gate")
    add_checklist_table(
        doc,
        ["Gate", "Criteria", "Met?"],
        [
            ["G1", "Critical admin flows A-01…A-18 and O-01…O-14 passed on staging", ""],
            ["G2", "No Sev-1 / Sev-2 open defects against release scope", ""],
            ["G3", "Security suite green on dedicated test DB (EC2 `hurix_talent_test` — done this cycle)", "Yes"],
            ["G4", "SMTP verified against production credentials", ""],
            ["G5", "Rollback owner named and contacted", ""],
            ["G6", "VP / Product written approval recorded below", ""],
        ],
    )

    add_body(
        doc,
        "Release decision options: GO  ·  CONDITIONAL GO (with listed waivers)  ·  NO-GO.",
        color=INK,
    )
    add_body(doc, "Decision (circle one):     GO          CONDITIONAL GO          NO-GO", color=SLATE)
    add_body(doc, "Conditions / waivers:", color=MUTED)
    # blank lines for handwriting
    for _ in range(3):
        p = doc.add_paragraph("_" * 92)
        clear_paragraph_spacing(p)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.color.rgb = RGBColor.from_string("CBD5E1")
            run.font.size = Pt(10)

    sp = doc.add_paragraph()
    clear_paragraph_spacing(sp)
    sp.paragraph_format.space_before = Pt(10)

    add_section_heading(doc, "12", "Approvals")
    # Remove accidental empty table from earlier helper mistake — only use add_blank_sign once
    table = doc.add_table(rows=4, cols=3)
    style_table_borders(table, color="CBD5E1", sz="4")
    for j, h in enumerate(["Role", "Name / Signature", "Date"]):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, HEADER_BG)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        clear_paragraph_spacing(p)
        add_runs(p, [(h, True, PAPER)], size=9)
    for i, label in enumerate(
        ["QA / Engineering Lead", "Product Owner", "Vice President (Go-Live Authority)"]
    ):
        for j, val in enumerate([label, "", ""]):
            cell = table.rows[i + 1].cells[j]
            set_cell_shading(cell, PAPER)
            set_cell_margins(cell, top=110, bottom=110, left=80, right=80)
            p = cell.paragraphs[0]
            clear_paragraph_spacing(p)
            add_runs(p, [(val, j == 0, SLATE)], size=9)

    footer = doc.add_paragraph()
    clear_paragraph_spacing(footer)
    footer.paragraph_format.space_before = Pt(16)
    add_runs(
        footer,
        [
            (
                "Hurix Digital  ·  Confidential  ·  Generated for pre-go-live QA review  ·  "
                f"{today}",
                False,
                MUTED,
            )
        ],
        size=8,
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
